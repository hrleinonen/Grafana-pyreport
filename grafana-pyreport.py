#!/usr/bin/env python3
#
# Grafana-pyreport version 0.92
# By: Ville Leinonen
# License: GPLv2
#
import argparse
import requests
import os
import os.path
from os import path
import datetime
from datetime import date, timedelta
from dateutil.relativedelta import *
import urllib3
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE

ver_number = "0.92"

# Get arguments from cli.
def arguments():

    parser = argparse.ArgumentParser(description="Grafana-pyreport " + str(ver_number) + " help.",
                                    formatter_class=argparse.ArgumentDefaultsHelpFormatter)
    parser.add_argument("-r", "--raw", action="store_true", help="Create raw report inc. image ID's.")
    parser.add_argument("-t", "--timeframe", help="Reporting timeframe 1d, 1w, 2w or 1m.", required=True)
    parser.add_argument("-l", "--template", help="Use template to create report.")
    parser.add_argument("-o", "--org", help="Organization (default is 1).")
    parser.add_argument("-z", "--timezone", help="Setup timezone to Grafana API call (default UTC, this should be safe choice).")
    parser.add_argument("-f", "--folder", help="Where to put generated reports (default is working dir).")
    parser.add_argument("apitoken", help="Grafana API token")
    parser.add_argument("desktop", help="Grafana desktop ID")
    parser.add_argument("url", help="Grafana server URL (eg. https://grafana.acme.local:3000)")
    args = parser.parse_args()
    config = vars(args)

    args = parser.parse_args()
    return args

# Check if running inside Docker container
def container_check():
    container_key = os.environ.get('DOCKER', False)

    if container_key:
        container = 1
    else:
        container = 0
    return container

# Generate time related arguments
def grafana_timeframe(time_frame):
    if time_frame.lower() == "1d":
        time_start = 1
    elif time_frame.lower() == "1w":
        time_start = 7
    elif time_frame.lower() == "2w":
        time_start = 14
    elif time_frame.lower() == "1m":
        time_start = 42
    else:
        print("Unknown error in timeframe.")
        exit()

    # Get current date and calculate prev dates.
    date_today = date.today()
    date_yesterday = date_today - timedelta(days = 1)
    if time_start == 42:
        date_start = date_yesterday - relativedelta(months = 1)
    else:
        date_start = date_yesterday - timedelta(days = time_start)

    # Add full hours etc.
    date_yesterday_sec = str(date_yesterday) + " 23:59:59"
    date_start_sec = str(date_start) + " 00:00:00"

    # Convert to Unix timestamp
    date_yesterday_convert = datetime.datetime.strptime(str(date_yesterday_sec),"%Y-%m-%d %H:%M:%S")
    date_yesterday_unix = datetime.datetime.timestamp(date_yesterday_convert)
    date_yesterday_unix_msec = date_yesterday_unix * 1e3
    date_yesterday_unix_msec = str(round(date_yesterday_unix_msec))

    date_start_convert = datetime.datetime.strptime(str(date_start_sec),"%Y-%m-%d %H:%M:%S")
    date_start_unix = datetime.datetime.timestamp(date_start_convert)
    date_start_unix_msec = date_start_unix * 1e3
    date_start_unix_msec = str(round(date_start_unix_msec))

    return date_today, date_yesterday, date_start, date_yesterday_unix_msec, date_start_unix_msec

# Get image id's from Grafana dashboard
def grafana_dashboard(apitoken, desktop, url):    

    # empty ID list
    id_list = []
    id_title = []
    id_row_list = []
    id_row_title = []

    headers = {"Authorization": "Bearer " + apitoken + ""}

    panel_id_url = url + "/api/dashboards/uid/" + desktop
    panel_id_data = requests.get(panel_id_url, headers=headers, verify=False)

    resp_dict = panel_id_data.json()
    resp_dict = resp_dict.get('dashboard')

    # Parse dashbord to values
    for key, val in resp_dict.items():
        if key == "title":
            dashboard_name = val
        if key == "panels":
            for pkey in val:
                id_check = 0 
                row_check = 0  
                title_check = 0
                for ikey, ival in pkey.items():
                    if ikey == "id":
                        id_check = 1
                        id = ival
                    if ikey == "type":
                        if ival == "row":
                            row_check = 1
                            row = ival
                    if ikey == "title":
                        title_check = 1
                        title = ival

                if row_check == 0:
                    id_list.append(id)
                    id_title.append(title)
                if row_check == 1:
                    id_row_list.append(id)
                    id_row_title.append(title)

    return dashboard_name, id_list, id_title, id_row_list, id_row_title

# Generate word report
def grafana_report(time_frame, time_zone, date_today, date_yesterday, date_start, date_yesterday_unix_msec, date_start_unix_msec,\
    raw, template, apitoken, desktop, url, dashboard_name, id_list, id_title, id_row_list, id_row_title, folder, org, container):

    headers = {"Authorization": "Bearer " + apitoken + ""}

    if not org:
        org = "1"

    if not time_zone:
        time_zone = "UTC"

    document = Document()

    # Document styles
    font = document.styles['Normal'].font
    font.name = 'Calibri'
    font.size = Pt(12)

    obj_styles = document.styles
    obj_charstyle = obj_styles.add_style('Big Style', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(20)
    obj_font.name = 'Calibri'

    obj_styles = document.styles
    obj_charstyle = obj_styles.add_style('Normal Style', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(12)
    obj_font.name = 'Calibri'

    # Front page
    paragraph_center = document.add_paragraph()
    paragraph_center.add_run("\n\n\n\n\n\n\n\n" + dashboard_name + "\n", style = 'Big Style')
    paragraph_center.add_run("Reporting period " + str(date_start) + " - " + str(date_yesterday) + "\n", style = 'Normal Style')
    paragraph_center.add_run("Report created " + str(date_today) + "", style = 'Normal Style')
    paragraph_format = paragraph_center.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_page_break()

    # Data pages
    if raw:
        for x in range(0, len(id_list)):
            print(".", end='', flush=True)
            image_url = str(url) + "/render/d-solo/" + desktop + "/" + dashboard_name + "?orgId=" + org + "&from=" + date_start_unix_msec + "&to="\
                + date_yesterday_unix_msec + "&panelId=" + str(id_list[x]) + "&width=1000&height=500&tz=" + str(time_zone)

            img_data = requests.get(image_url, headers=headers, verify=False, timeout=120).content
            with open("/tmp/grafana_image_" + str(id_list[x]) + ".jpg", "wb") as handler:
                handler.write(img_data)
            document.add_picture("/tmp/grafana_image_" + str(id_list[x]) + ".jpg", width=Cm(12))
            os.remove("/tmp/grafana_image_" + str(id_list[x]) + ".jpg")

            p = document.add_paragraph("<id:" + str(id_list[x]) + ">, " + str(id_title[x]))
    
    elif not template:
        for x in range(0, len(id_list)):
            print(".", end='', flush=True)
            image_url = str(url) + "/render/d-solo/" + desktop + "/" + dashboard_name + "?orgId=" + org + "&from=" + date_start_unix_msec + "&to="\
                 + date_yesterday_unix_msec + "&panelId=" + str(id_list[x]) + "&width=1000&height=500&tz=" + str(time_zone)

            img_data = requests.get(image_url, headers=headers, verify=False, timeout=120).content
            with open("/tmp/grafana_image_" + str(id_list[x]) + ".jpg", "wb") as handler:
                handler.write(img_data)
            document.add_picture("/tmp/grafana_image_" + str(id_list[x]) + ".jpg", width=Cm(12))
            os.remove("/tmp/grafana_image_" + str(id_list[x]) + ".jpg")

            paragraph = document.add_paragraph()
            paragraph.style = 'List Number'
            paragraph.add_run(str(id_title[x]) + "\n", style = 'Normal Style')

    else:
        if container:
            template_file = open("/opt/grafana-pyreport/templates/" + template, 'r')
        else:
            template_file = open(template, 'r')
        Lines = template_file.readlines()

        for line in Lines:
            print(".", end='', flush=True)
            if "<title>" in line:
                line = line.strip("<title>").strip()
                document.add_heading(str(line))
            elif "<p>" in line:
                document.add_page_break()
            elif "<id:" in line:
                line = line.strip("<:id").strip()
                line = line.strip(">").strip()

                image_url = str(url) + "/render/d-solo/" + desktop + "/" + dashboard_name + "?orgId=" + org + "&from=" + date_start_unix_msec + "&to="\
                     + date_yesterday_unix_msec + "&panelId=" + str(line) + "&width=1000&height=500&tz=" + str(time_zone)

                # Check that image found in Grafana using id_list
                if int(line) in id_list:
                    img_data = requests.get(image_url, headers=headers, verify=False, timeout=120).content
                    with open("/tmp/grafana_image_" + str(line) + ".jpg", "wb") as handler:
                        handler.write(img_data)
                    document.add_picture("/tmp/grafana_image_" + str(line) + ".jpg", width=Cm(12))
                    os.remove("/tmp/grafana_image_" + str(line) + ".jpg")

                    id_index = id_list.index(int(line))
                    paragraph = document.add_paragraph()
                    paragraph.style = 'List Number'
                    paragraph.add_run(str(id_title[id_index]), style = 'Normal Style')
                else:
                    paragraph = document.add_paragraph()
                    paragraph.style = 'List Number'
                    paragraph.add_run("Image id " + line + " not found in Grafana dashboard or rendering takes over 120 second.", style = 'Normal Style') 
                 
            else:
                line = line.strip()
                paragraph = document.add_paragraph()
                paragraph.add_run(str(line), style = 'Normal Style')

    # Output to Word file
    if raw:
        if folder:
            if container:
                document.save("/opt/grafana-pyreport/output/" + folder + "/" + dashboard_name + "-raw-" + str(date_start) + "-" + str(date_yesterday) + ".docx")
            else:
                document.save(folder + "/" + dashboard_name + "-raw-" + str(date_start) + "-" + str(date_yesterday) + ".docx")
        else:
            if container:
                document.save("/opt/grafana-pyreport/output/" + dashboard_name + "-raw-" + str(date_start) + "-" + str(date_yesterday) + ".docx")
            else:
                document.save(dashboard_name + "-raw-" + str(date_start) + "-" + str(date_yesterday) + ".docx")
        document_name = dashboard_name + "-raw-" + str(date_start) + "-" + str(date_yesterday) + ".docx"
    elif not template:
        if folder:
            if container:
                document.save("/opt/grafana-pyreport/output/" + folder + "/" + dashboard_name + "-prensentation-" + str(date_start) + "-" + str(date_yesterday) + ".docx")
            else:
                document.save(folder + "/" + dashboard_name + "-prensentation-" + str(date_start) + "-" + str(date_yesterday) + ".docx")
        else:
            if container:
                document.save("/opt/grafana-pyreport/output/" + dashboard_name + "-prensentation-" + str(date_start) + "-" + str(date_yesterday) + ".docx")
            else:
                document.save(dashboard_name + "-prensentation-" + str(date_start) + "-" + str(date_yesterday) + ".docx")
        document_name = dashboard_name + "-prensentation-" + str(date_start) + "-" + str(date_yesterday) + ".docx"        
    else:
        if folder:
            if container:
                document.save("/opt/grafana-pyreport/output/" + folder + "/" + dashboard_name + "_" + str(date_start) + "-" + str(date_yesterday) + ".docx")
            else:
                document.save(folder + "/" + dashboard_name + "_" + str(date_start) + "-" + str(date_yesterday) + ".docx")
        else:
            if container:
                document.save("/opt/grafana-pyreport/output/" + dashboard_name + "_" + str(date_start) + "-" + str(date_yesterday) + ".docx")
            else:
                document.save(dashboard_name + "_" + str(date_start) + "-" + str(date_yesterday) + ".docx")
        document_name = dashboard_name + "_" + str(date_start) + "-" + str(date_yesterday) + ".docx"

    return document_name

# Main function
def main():

    # Parse commandline argiments
    args = arguments()
    time_frame = args.timeframe
    time_zone = args.timezone
    raw = args.raw
    template = args.template
    apitoken = args.apitoken
    desktop = args.desktop
    url = args.url
    folder = args.folder
    org = args.org

    # Check if inside Docker container
    container = container_check()

    if raw and template:
        print("Please do not use presentation and template switches together.")
        exit()
    
    if container:
        if template and not path.exists("/opt/grafana-pyreport/templates/" + template):
            print("Template file {} not found.".format(template))
            exit()
    else:
        if template and not path.exists(template):
            print("Template file {} not found.".format(template))
            exit()

    print("Starting generating report using timeframe {}".format(time_frame))
    
    # Create reporting timeframe
    date_today, date_yesterday, date_start, date_yesterday_unix_msec, date_start_unix_msec = grafana_timeframe(time_frame)

    # Get panel and row ids from dashboard
    dashboard_name, id_list, id_title, id_row_list, id_row_title = grafana_dashboard(apitoken, desktop, url)

    # Create word report
    print("Getting data from Grafana {}, this can take long time.".format(url))
    document_name = grafana_report(time_frame, time_zone, date_today, date_yesterday, date_start, date_yesterday_unix_msec, date_start_unix_msec,\
         raw, template, apitoken, desktop, url, dashboard_name, id_list, id_title, id_row_list, id_row_title, folder, org, container)
    
    print("\nReporting finished, report name is {}".format(str(document_name)))

if __name__ == "__main__":
    main()
